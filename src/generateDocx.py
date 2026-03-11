import csv
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, RGBcolor, EMULTIHOP
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from ReportConfig import colors, font_office, report_title, sizes, report_subtitle, report_subtitle, Sections

# PATH
BASE_DIR = Path(__file__).resolve().parent.parent
REPORT_DATA_DIR = Path(__file__).resolve().parent/"report_data"
OUTPUT_DIR = BASE_DIR/"outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

#Helpers

def hex_to_rgb(hex_color):
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2]), 16), int(h[2:4], 16), int(h[4:6], 16)

def load_csv_data(filename):
    path = REPORT_DATA_DIR/filename
    with open(path, encoding="utf-8") as f:
        return list(csv.DictReader(f))

def fmt_dollar(value):
    return f"${float(value):,.2f}"

def fmt_number(value):
    return f"{int(float(value)):,}"

def fmt_pct(value):
    v = float(value)
    return f"{'+' if v >= 0 else ''}{v}%"

PRIMARY = hex_to_rgb(colors["primary"])
SECONDARY = hex_to_rgb(colors["secondary"])
ACCENT = hex_to_rgb(colors["accent"])
SUCCESS = hex_to_rgb(colors["success"])
DARK_TEXT = hex_to_rgb(colors["dark_text"])
FONT = font_office["heading"]

def set_cell_shading(cell, hex_color):
    color = hex_color.lstrip("#")
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:color="{val.get("color", "BDC3C7")}" '
            f'w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths, highlight_col=None, highlight_colors=None):
    """Create a formatted table with header row and alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    for i, (header, width) in enumerate(zip(headers, col_widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(width)
        set_cell_shading(cell, colors["primary"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(sizes["table_head"])
        run.font.color.rgb = hex_to_rgb(colors["light_text"])
        run.font.name = FONT

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.width = Inches(col_widths[col_idx])

            # Alternate row shading
            if row_idx % 2 == 0:
                set_cell_shading(cell, colors["light_bg"])

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            run.font.size = Pt(sizes["table_body"])
            run.font.name = FONT
            run.font.color.rgb = DARK_TEXT

            # Highlight column with color
            if highlight_col is not None and col_idx == highlight_col and highlight_colors:
                try:
                    val = float(str(value).replace("%", "").replace("+", "").replace(",", ""))
                    run.bold = True
                    run.font.color.rgb = highlight_colors[0] if val >= 0 else highlight_colors[1]
                except ValueError:
                    pass

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:color="{colors["border"].lstrip("#")}"/>'
        f'  <w:left w:val="single" w:sz="4" w:color="{colors["border"].lstrip("#")}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{colors["border"].lstrip("#")}"/>'
        f'  <w:right w:val="single" w:sz="4" w:color="{colors["border"].lstrip("#")}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:color="{colors["border"].lstrip("#")}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:color="{colors["border"].lstrip("#")}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    return table


def add_section_heading(doc, text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.color.rgb = PRIMARY
    run.font.name = FONT
    run.font.size = Pt(sizes["section"])
    # Bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:color="{colors["primary"].lstrip("#")}" w:space="1"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_body_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(sizes["body"])
    run.font.name = FONT
    run.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after = Pt(8)
    return p


def add_subtitle_label(doc, text, color):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes["body"])
    run.font.name = FONT
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


# ──────────────────────────────────────────────
# KPI CARD ROW
# ──────────────────────────────────────────────

def add_kpi_row(doc, kpis):
    """Add a row of KPI cards as a borderless table."""
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_width = 6.5 / len(kpis)

    for i, kpi in enumerate(kpis):
        # Value row
        cell_val = table.rows[0].cells[i]
        cell_val.width = Inches(col_width)
        if i % 2 == 0:
            set_cell_shading(cell_val, colors["light_bg"])
        p = cell_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(kpi["value"])
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = FONT
        run.font.color.rgb = PRIMARY

        # Label row
        cell_lbl = table.rows[1].cells[i]
        cell_lbl.width = Inches(col_width)
        if i % 2 == 0:
            set_cell_shading(cell_lbl, colors["light_bg"])
        p2 = cell_lbl.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(kpi["label"])
        run2.font.size = Pt(8)
        run2.font.name = FONT
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Minimal borders (top/bottom only)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:color="{colors["primary"].lstrip("#")}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:color="{colors["primary"].lstrip("#")}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    return table


# ──────────────────────────────────────────────
# HEADER / FOOTER
# ──────────────────────────────────────────────

def add_header_footer(doc):
    section = doc.sections[0]

    # Header
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_left = header_p.add_run(report_title)
    run_left.bold = True
    run_left.font.size = Pt(8)
    run_left.font.name = FONT
    run_left.font.color.rgb = PRIMARY

    run_tab = header_p.add_run("\t\t\t\t\t\t\t")
    run_right = header_p.add_run("Fiscal Year 2016")
    run_right.font.size = Pt(8)
    run_right.font.name = FONT
    run_right.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Bottom border on header
    pPr = header_p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{colors["secondary"].lstrip("#")}" w:space="1"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer_p = footer.paragraphs[0]
    run_conf = footer_p.add_run("Confidential — Internal Use Only")
    run_conf.font.size = Pt(sizes["footer"])
    run_conf.font.name = FONT
    run_conf.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ──────────────────────────────────────────────
# MAIN GENERATOR
# ──────────────────────────────────────────────

def generate_report():
    print("Loading processed data...")
    es = load_csv_data("executive_summary.csv")[0]
    stores = load_csv_data("store_overview.csv")
    top_margin = load_csv_data("top_margin.csv")
    bottom_margin = load_csv_data("bottom_margin.csv")
    fastest = load_csv_data("fastest_moving.csv")
    slowest = load_csv_data("slowest_moving.csv")
    vendors = load_csv_data("vendor_top10.csv")

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Default font ──
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(sizes["body"])
    style.font.color.rgb = DARK_TEXT

    add_header_footer(doc)

    # ═══════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════

    doc.add_paragraph()  # spacer

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(report_title)
    run.bold = True
    run.font.size = Pt(sizes["title"])
    run.font.name = FONT
    run.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run(report_subtitle)
    run2.font.size = Pt(sizes["subtitle"] - 3)
    run2.font.name = FONT
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # ═══════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════

    add_section_heading(doc, Sections[0])

    add_kpi_row(doc, [
        {"label": "Stores", "value": es["total_stores"]},
        {"label": "Active SKUs (EOY)", "value": fmt_number(es["unique_skus_end"])},
        {"label": "Avg Margin", "value": f"{es['avg_margin_pct']}%"},
        {"label": "Vendors", "value": es["total_vendors"]},
    ])

    doc.add_paragraph()  # spacer

    summary_text = (
        f"During fiscal year 2016, total inventory units grew from "
        f"{fmt_number(es['beg_total_units'])} to {fmt_number(es['end_total_units'])} "
        f"({fmt_pct(es['unit_change_pct'])}), representing an inventory value increase from "
        f"{fmt_dollar(es['beg_total_value'])} to {fmt_dollar(es['end_total_value'])} "
        f"({fmt_pct(es['value_change_pct'])}). The product catalog expanded from "
        f"{fmt_number(es['unique_skus_beg'])} to {fmt_number(es['unique_skus_end'])} SKUs. "
        f"Total procurement spend reached {fmt_dollar(es['total_purchase_spend'])} with "
        f"{fmt_dollar(es['total_freight'])} in freight costs across {es['total_vendors']} "
        f"active suppliers. The average product margin stands at {es['avg_margin_pct']}%, "
        f"and {es['dead_stock_count']} products showed zero movement throughout the year."
    )
    add_body_text(doc, summary_text)

    # ═══════════════════════════════════════════
    # 2. INVENTORY OVERVIEW
    # ═══════════════════════════════════════════

    doc.add_page_break()
    add_section_heading(doc, Sections[1])
    add_body_text(doc, "Top 10 stores by end-of-year inventory value:")

    store_headers = ["Store", "City", "Beg Units", "End Units", "Delta", "Beg Value", "End Value"]
    store_widths = [0.6, 1.2, 0.9, 0.9, 0.8, 1.1, 1.0]
    store_rows = []
    for s in stores:
        delta = int(float(s["delta_units"]))
        store_rows.append([
            s["store"], s["city"],
            fmt_number(s["beg_units"]), fmt_number(s["end_units"]),
            f"{'+' if delta >= 0 else ''}{delta:,}",
            fmt_dollar(s["beg_value"]), fmt_dollar(s["end_value"]),
        ])

    add_styled_table(doc, store_headers, store_rows, store_widths,
                     highlight_col=4, highlight_colors=[ACCENT, SUCCESS])

    doc.add_paragraph()
    top_store = stores[0]
    add_body_text(doc, (
        f"The store with the highest inventory value at year-end is Store {top_store['store']} "
        f"in {top_store['city']}, holding {fmt_number(top_store['end_units'])} units valued at "
        f"{fmt_dollar(top_store['end_value'])}. Across all stores, inventory grew by "
        f"{fmt_pct(es['unit_change_pct'])} in units and {fmt_pct(es['value_change_pct'])} in value, "
        f"indicating significant stock accumulation that may require attention regarding "
        f"storage capacity and carrying costs."
    ))

    # ═══════════════════════════════════════════
    # 3. MARGIN ANALYSIS
    # ═══════════════════════════════════════════

    doc.add_page_break()
    add_section_heading(doc, Sections[2])
    add_body_text(doc, (
        f"The average margin across all products is {es['avg_margin_pct']}%. "
        f"Below are the top and bottom 10 products by margin percentage."
    ))

    margin_headers = ["Product", "Sale Price", "Cost", "Margin $", "Margin %", "Brand"]
    margin_widths = [2.1, 0.8, 0.8, 0.8, 0.7, 0.6]

    def margin_to_rows(items):
        rows = []
        for m in items:
            rows.append([
                m["description"][:40],
                fmt_dollar(m["price"]), fmt_dollar(m["cost"]),
                fmt_dollar(m["margin_abs"]), f"{m['margin_pct']}%",
                m["brand"],
            ])
        return rows

    add_subtitle_label(doc, "Top 10 — Highest Margin Products", SUCCESS)
    add_styled_table(doc, margin_headers, margin_to_rows(top_margin), margin_widths,
                     highlight_col=4, highlight_colors=[SUCCESS, ACCENT])

    add_subtitle_label(doc, "Bottom 10 — Lowest Margin Products", ACCENT)
    add_styled_table(doc, margin_headers, margin_to_rows(bottom_margin), margin_widths,
                     highlight_col=4, highlight_colors=[SUCCESS, ACCENT])

    # ═══════════════════════════════════════════
    # 4. STOCK ROTATION
    # ═══════════════════════════════════════════

    doc.add_page_break()
    add_section_heading(doc, Sections[3])
    add_body_text(doc, (
        f"Stock rotation analysis compares beginning and end-of-year inventory levels per product. "
        f"Products with the largest decrease in stock are considered fast-moving (high demand), "
        f"while those with the largest increase indicate slow rotation or over-purchasing. "
        f"A total of {es['dead_stock_count']} products showed zero movement during the year (dead stock)."
    ))

    rot_headers = ["Product", "Beg Stock", "End Stock", "Delta", "Brand"]
    rot_widths = [2.4, 1.0, 1.0, 1.0, 0.8]

    def rotation_to_rows(items):
        rows = []
        for r in items:
            delta = int(float(r["delta"]))
            rows.append([
                r["description"][:40],
                fmt_number(r["beg_stock"]), fmt_number(r["end_stock"]),
                f"{'+' if delta >= 0 else ''}{delta:,}",
                r["brand"],
            ])
        return rows

    add_subtitle_label(doc, "Top 10 — Fastest Moving (biggest stock decrease)", SUCCESS)
    add_styled_table(doc, rot_headers, rotation_to_rows(fastest), rot_widths,
                     highlight_col=3, highlight_colors=[ACCENT, SUCCESS])

    add_subtitle_label(doc, "Top 10 — Slowest Moving (biggest stock increase)", ACCENT)
    add_styled_table(doc, rot_headers, rotation_to_rows(slowest), rot_widths,
                     highlight_col=3, highlight_colors=[ACCENT, SUCCESS])

    # ═══════════════════════════════════════════
    # 5. SUPPLIER ANALYSIS
    # ═══════════════════════════════════════════

    doc.add_page_break()
    add_section_heading(doc, Sections[4])
    add_body_text(doc, (
        f"The company works with {es['total_vendors']} active suppliers. "
        f"Total procurement spend in 2016 was {fmt_dollar(es['total_purchase_spend'])} "
        f"plus {fmt_dollar(es['total_freight'])} in freight. "
        f"Below are the top 10 suppliers by total spend."
    ))

    vendor_headers = ["Supplier", "Total Spend", "Freight", "Orders", "Avg Pay Days", "% Total"]
    vendor_widths = [1.8, 1.3, 0.9, 0.7, 0.8, 0.7]
    vendor_rows = []
    for v in vendors:
        vendor_rows.append([
            v["name"],
            fmt_dollar(v["spend"]), fmt_dollar(v["freight"]),
            v["orders"], f"{v['avg_pay_days']} d",
            f"{v['pct_of_total']}%",
        ])

    add_styled_table(doc, vendor_headers, vendor_rows, vendor_widths,
                     highlight_col=5, highlight_colors=[PRIMARY, PRIMARY])

    doc.add_paragraph()
    top_v = vendors[0]
    add_body_text(doc, (
        f"The largest supplier is {top_v['name'].strip()}, accounting for "
        f"{top_v['pct_of_total']}% of total procurement spend ({fmt_dollar(top_v['spend'])}). "
        f"Average payment terms across all suppliers range from "
        f"{min(int(v['avg_pay_days']) for v in vendors)} to "
        f"{max(int(v['avg_pay_days']) for v in vendors)} days."
    ))

    # ── Save ──
    output_path = OUTPUT_DIR / "WarehouseReport.docx"
    doc.save(str(output_path))
    print(f"\n✓ Report saved to {output_path}")


if __name__ == "__main__":
    generate_report()
